#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
办公自动化工具 - Web版 (Flask)
Excel收集表批量导入 + Word模板导出
"""

import os
import logging
import configparser
import tempfile
from io import BytesIO
from flask import Flask, render_template, request, send_file, redirect, url_for, flash, jsonify
from werkzeug.utils import secure_filename
import pandas as pd
from docx import Document

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.urandom(24)

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal Server Error: {e}")
    return """<!DOCTYPE html><html><head><title>Error</title>
    <style>body{font-family:sans-serif;display:flex;justify-content:center;align-items:center;height:100vh;background:#f0f2f5}
    .box{background:white;padding:40px;border-radius:12px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.1)}
    a{color:#667eea;text-decoration:none}</style></head>
    <body><div class="box"><h1>😵 出错了</h1><p>服务器内部错误，请刷新重试或检查配置。</p><a href="/">← 返回首页</a></div></body></html>""", 500

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 配置管理
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CONFIG_FILE = os.environ.get("CONFIG_FILE", "/data/config.ini")
DATA_DIR = os.environ.get("DATA_DIR", "/data")


def load_config() -> dict:
    cfg = configparser.ConfigParser()
    if os.path.exists(CONFIG_FILE):
        cfg.read(CONFIG_FILE, encoding="utf-8")
    return {
        "master_path": cfg.get("paths", "master_excel", fallback=""),
        "template_path": cfg.get("paths", "word_template", fallback=""),
    }


def save_config(master_path: str, template_path: str):
    cfg = configparser.ConfigParser()
    cfg["paths"] = {
        "master_excel": master_path,
        "word_template": template_path,
    }
    os.makedirs(os.path.dirname(CONFIG_FILE), exist_ok=True)
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        cfg.write(f)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Excel 导入逻辑
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def read_master(master_path: str) -> pd.DataFrame:
    if not os.path.isfile(master_path):
        return pd.DataFrame()
    return pd.read_excel(master_path, engine="openpyxl", dtype=str)


def read_collection(file_path: str) -> pd.DataFrame:
    df = pd.read_excel(file_path, engine="openpyxl", header=None, dtype=str)
    return df


def validate_collection(coll_df: pd.DataFrame, master_headers: list) -> tuple:
    coll_headers = coll_df.iloc[:, 0].tolist()
    if len(coll_headers) != len(master_headers):
        missing = set(master_headers) - set(coll_headers)
        extra = set(coll_headers) - set(master_headers)
        parts = []
        if missing:
            parts.append(f"缺少字段: {', '.join(missing)}")
        if extra:
            parts.append(f"多余字段: {', '.join(extra)}")
        return False, "表头不一致！\n" + "\n".join(parts)

    diff_fields = []
    for i, (mh, ch) in enumerate(zip(master_headers, coll_headers)):
        if str(mh).strip() != str(ch).strip():
            diff_fields.append(f" 第{i+1}行: 总表[{mh}] ≠ 收集表[{ch}]")
    if diff_fields:
        return False, "表头不一致！\n" + "\n".join(diff_fields)

    values = coll_df.iloc[:, 1].tolist()
    empty_fields = []
    for header, val in zip(coll_headers, values):
        if pd.isna(val) or str(val).strip() == "":
            empty_fields.append(f" - {header}")
    if empty_fields:
        return False, "以下字段未填写：\n" + "\n".join(empty_fields)

    return True, ""


def import_collection(master_df: pd.DataFrame, coll_df: pd.DataFrame) -> pd.DataFrame:
    headers = coll_df.iloc[:, 0].tolist()
    values = coll_df.iloc[:, 1].tolist()
    pk_field = headers[0]
    pk_value = str(values[0]).strip()

    row_data = {h: str(v).strip() if not pd.isna(v) else "" for h, v in zip(headers, values)}

    if master_df.empty:
        master_df = pd.DataFrame(columns=headers)

    if pk_field in master_df.columns:
        mask = master_df[pk_field].astype(str).str.strip() == pk_value
    else:
        mask = pd.Series([False] * len(master_df))

    if mask.any():
        for col, val in row_data.items():
            if col in master_df.columns:
                master_df.loc[mask, col] = val
    else:
        new_row = pd.DataFrame([row_data])
        master_df = pd.concat([master_df, new_row], ignore_index=True)

    return master_df


def save_master(master_df: pd.DataFrame, master_path: str):
    os.makedirs(os.path.dirname(master_path), exist_ok=True)
    master_df.to_excel(master_path, index=False, engine="openpyxl")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Word 导出逻辑
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def replace_placeholder(text: str, replacements: dict) -> str:
    for key, val in replacements.items():
        text = text.replace(f"{{{{{key}}}}}", str(val))
    return text


def export_word(template_path: str, row_data: dict, save_path: str):
    doc = Document(template_path)

    for para in doc.paragraphs:
        if "{{" in para.text:
            full_text = para.text
            new_text = replace_placeholder(full_text, row_data)
            if new_text != full_text:
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text:
                        full_text = para.text
                        new_text = replace_placeholder(full_text, row_data)
                        if new_text != full_text:
                            for i, run in enumerate(para.runs):
                                if i == 0:
                                    run.text = new_text
                                else:
                                    run.text = ""

    doc.save(save_path)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Flask 路由
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.route("/")
def index():
    config = load_config()
    return render_template("index.html", config=config)


@app.route("/config", methods=["GET", "POST"])
def config_page():
    if request.method == "POST":
        master_path = request.form.get("master_path", "").strip()
        template_path = request.form.get("template_path", "").strip()
        if not master_path or not template_path:
            flash("两个路径都不能为空！", "error")
            return redirect(url_for("config_page"))
        save_config(master_path, template_path)
        flash("配置已保存！", "success")
        return redirect(url_for("index"))
    config = load_config()
    return render_template("config.html", config=config)


@app.route("/import", methods=["GET", "POST"])
def import_page():
    config = load_config()
    master_path = config.get("master_path", "")

    if request.method == "POST":
        if not master_path:
            flash("请先配置总表 Excel 路径！", "error")
            return redirect(url_for("config_page"))

        files = request.files.getlist("files")
        if not files or all(f.filename == "" for f in files):
            flash("请选择至少一个收集表文件！", "error")
            return redirect(url_for("import_page"))

        try:
            master_df = read_master(master_path)
        except Exception as e:
            flash(f"读取总表失败：{e}", "error")
            return redirect(url_for("import_page"))

        master_headers = list(master_df.columns) if not master_df.empty else None

        success_count = 0
        fail_count = 0
        fail_details = []

        for f in files:
            if not f.filename:
                continue
            fname = secure_filename(f.filename)
            if not fname:
                fname = f.filename

            try:
                with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                    f.save(tmp.name)
                    coll_df = read_collection(tmp.name)
                os.unlink(tmp.name)
            except Exception as e:
                fail_count += 1
                fail_details.append(f"【{fname}】读取失败：{e}")
                continue

            if master_headers is None:
                master_headers = coll_df.iloc[:, 0].tolist()

            passed, msg = validate_collection(coll_df, master_headers)
            if not passed:
                fail_count += 1
                fail_details.append(f"【{fname}】验证不通过：\n{msg}")
                continue

            try:
                master_df = import_collection(master_df, coll_df)
                success_count += 1
            except Exception as e:
                fail_count += 1
                fail_details.append(f"【{fname}】导入异常：{e}")

        if success_count > 0:
            try:
                save_master(master_df, master_path)
            except Exception as e:
                flash(f"保存总表失败：{e}", "error")
                return redirect(url_for("import_page"))

        summary = f"导入完成！成功：{success_count} 个，失败：{fail_count} 个"
        if fail_details:
            summary += " | 失败详情：" + "；".join(fail_details)

        flash(summary, "success" if fail_count == 0 else "warning")
        return redirect(url_for("import_page"))

    return render_template("import.html", config=config)


@app.route("/import_template", methods=["GET", "POST"])
def import_template_page():
    """上传新的 Word 模板，覆盖配置中的模板文件"""
    config = load_config()
    template_path = config.get("template_path", "")

    if request.method == "POST":
        if not template_path:
            flash("请先在配置页面设置 Word 模板路径！", "error")
            return redirect(url_for("config_page"))

        f = request.files.get("template_file")
        if not f or f.filename == "":
            flash("请选择要上传的 Word 模板文件！", "error")
            return redirect(url_for("import_template_page"))

        fname = secure_filename(f.filename) or f.filename
        if not fname.lower().endswith(".docx"):
            flash("请上传 .docx 格式的 Word 文件！", "error")
            return redirect(url_for("import_template_page"))

        tmp_path = None
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
            os.close(tmp_fd)
            f.save(tmp_path)
            # 验证文件能正常打开
            doc = Document(tmp_path)
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            # 确保目标目录存在
            template_dir = os.path.dirname(template_path)
            if template_dir:
                os.makedirs(template_dir, exist_ok=True)
            # 覆盖原模板
            with open(tmp_path, "rb") as src:
                with open(template_path, "wb") as dst:
                    dst.write(src.read())
        except Exception as e:
            flash(f"上传失败：{e}", "error")
            return redirect(url_for("import_template_page"))
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        flash(f"Word 模板已更新！段落数：{para_count}，表格数：{table_count}", "success")
        return redirect(url_for("import_template_page"))

    # 检查当前模板是否存在
    template_exists = template_path and os.path.isfile(template_path)
    template_info = None
    if template_exists:
        try:
            doc = Document(template_path)
            placeholders = set()
            for para in doc.paragraphs:
                import re
                placeholders.update(re.findall(r"\{\{(.+?)\}\}", para.text))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            placeholders.update(re.findall(r"\{\{(.+?)\}\}", para.text))
            template_info = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "placeholders": sorted(placeholders),
            }
        except Exception:
            template_info = None

    return render_template("import_template.html", config=config, template_exists=template_exists, template_info=template_info)


@app.route("/export", methods=["GET", "POST"])
def export_page():
    config = load_config()
    master_path = config.get("master_path", "")
    template_path = config.get("template_path", "")

    if request.method == "POST":
        selected = request.form.get("pk_value", "").strip()
        if not selected:
            flash("请选择一个编号！", "error")
            return redirect(url_for("export_page"))

        try:
            master_df = read_master(master_path)
        except Exception as e:
            flash(f"读取总表失败：{e}", "error")
            return redirect(url_for("export_page"))

        pk_col = master_df.columns[0]
        mask = master_df[pk_col].astype(str).str.strip() == selected
        if not mask.any():
            flash(f"未找到编号 [{selected}] 对应的数据。", "error")
            return redirect(url_for("export_page"))

        row = master_df[mask].iloc[0].to_dict()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            export_word(template_path, row, tmp.name)
            buf = BytesIO()
            with open(tmp.name, "rb") as f:
                buf.write(f.read())
            buf.seek(0)
            os.unlink(tmp.name)

        return send_file(
            buf,
            as_attachment=True,
            download_name=f"{selected}_导出.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # GET: 读取总表获取编号列表
    pk_values = []
    if master_path and os.path.isfile(master_path):
        try:
            master_df = read_master(master_path)
            if not master_df.empty:
                pk_col = master_df.columns[0]
                pk_values = master_df[pk_col].dropna().astype(str).str.strip().tolist()
        except Exception:
            pass

    return render_template("export.html", config=config, pk_values=pk_values)


@app.route("/preview")
def preview_page():
    """预览总表（只读）"""
    config = load_config()
    master_path = config.get("master_path", "")
    headers = []
    rows = []
    row_count = 0

    if master_path and os.path.isfile(master_path):
        try:
            master_df = read_master(master_path)
            if not master_df.empty:
                headers = list(master_df.columns)
                rows = master_df.fillna("").values.tolist()
                row_count = len(rows)
        except Exception as e:
            flash(f"读取总表失败：{e}", "error")

    return render_template("preview.html", config=config, headers=headers, rows=rows, row_count=row_count)


@app.route("/download_master")
def download_master():
    """导出总表 Excel 文件"""
    config = load_config()
    master_path = config.get("master_path", "")
    if not master_path or not os.path.isfile(master_path):
        flash("总表文件不存在！", "error")
        return redirect(url_for("index"))

    buf = BytesIO()
    with open(master_path, "rb") as f:
        buf.write(f.read())
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="总表.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@app.route("/overwrite_master", methods=["POST"])
def overwrite_master():
    """上传修改后的总表覆盖原文件"""
    config = load_config()
    master_path = config.get("master_path", "")
    if not master_path:
        flash("请先配置总表路径！", "error")
        return redirect(url_for("config_page"))

    f = request.files.get("master_file")
    if not f or f.filename == "":
        flash("请选择要上传的总表文件！", "error")
        return redirect(url_for("preview_page"))

    tmp_path = None
    try:
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".xlsx")
        os.close(tmp_fd)
        f.save(tmp_path)
        # 验证文件能正常读取
        test_df = pd.read_excel(tmp_path, engine="openpyxl", dtype=str)
        if test_df.empty:
            flash("上传的文件为空！", "error")
            return redirect(url_for("preview_page"))
        # 确保目录存在
        master_dir = os.path.dirname(master_path)
        if master_dir:
            os.makedirs(master_dir, exist_ok=True)
        # 覆盖原文件
        with open(tmp_path, "rb") as src:
            with open(master_path, "wb") as dst:
                dst.write(src.read())
        row_count = len(test_df)
    except Exception as e:
        flash(f"覆盖总表失败：{e}", "error")
        return redirect(url_for("preview_page"))
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    flash(f"总表已成功覆盖！共 {row_count} 行数据。", "success")
    return redirect(url_for("preview_page"))


@app.route("/api/pk_values")
def api_pk_values():
    """API: 获取总表中的编号列表（供搜索用）"""
    config = load_config()
    master_path = config.get("master_path", "")
    if not master_path or not os.path.isfile(master_path):
        return jsonify([])
    try:
        master_df = read_master(master_path)
        if master_df.empty:
            return jsonify([])
        pk_col = master_df.columns[0]
        values = master_df[pk_col].dropna().astype(str).str.strip().unique().tolist()
        return jsonify(sorted(values))
    except Exception:
        return jsonify([])


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 启动
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
